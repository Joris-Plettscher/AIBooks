#author: Joris Plettscher
#Please read ChatGPT's policies before use
from selenium import webdriver
from selenium.webdriver.support import expected_conditions as expect
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
import undetected_chromedriver as uc
import random
import pyperclip
from docx import Document
import pyautogui
import time
import atexit

#Your ChatGPT login-data:
email_address = ""
password = ""

genre_f = open("genres.txt", "r") #File containing a list of book genres - foramt: genre1,genre2,...
#Choosing random genre
genres = genre_f.read().split(",")
genre = genres[random.randint(0,len(genres)-1)]

driver = uc.Chrome() #Undetected chromedriver as webdriver

#Quitting driver on exit
def exit_handler():
    driver.quit()

atexit.register(exit_handler)

document = Document() #Word document containing the resulting book

driver.get("https://chat.openai.com/auth/login")
driver.maximize_window()

pyperclip.copy('')

#XPATHs of used elements (Modify if necessary)
loginX = "//button[@class='btn relative btn-primary'][.='Log in']" #Button to navigate to the login page
nextX = "//button[@class='btn relative btn-neutral ml-auto']" #Button to skip the tutorial
doneX = "//button[@class='btn relative btn-primary ml-auto']" #Button to finish the tutorial
messageX = "//textarea[@class='m-0 w-full resize-none border-0 bg-transparent p-0 pr-10 focus:ring-0 focus-visible:ring-0 dark:bg-transparent md:pr-12 pl-3 md:pl-0']" #Textfield for messages to ChatGPT
sendX = "//button[@class='absolute p-1 rounded-md md:bottom-3 md:p-2 md:right-3 dark:hover:bg-gray-900 dark:disabled:hover:bg-transparent right-2 disabled:text-gray-400 enabled:bg-brand-purple text-white bottom-1.5 transition-colors disabled:opacity-40']" #Button to send messages
copyX = "(//button[@class='flex ml-auto gap-2 rounded-md p-1 hover:bg-gray-100 hover:text-gray-700 dark:text-gray-400 dark:hover:bg-gray-700 dark:hover:text-gray-200 disabled:dark:hover:text-gray-400'])" #Button to copy ChatGPT's response
scrollX = "(//button[@class='cursor-pointer absolute right-6 bottom-[124px] md:bottom-[120px] z-10 rounded-full border border-gray-200 bg-gray-50 text-gray-600 dark:border-white/10 dark:bg-white/10 dark:text-gray-200'])" #Button to scroll down to the bottom of the messages

#Login
login = WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.XPATH, loginX)))
ActionChains(driver).click(login).perform()

email = WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.ID, "username")))
email.send_keys(email_address)

email.send_keys(Keys.RETURN)

pwd =  WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.ID, "password")))
pwd.send_keys(password)

pwd.send_keys(Keys.RETURN)

#Skipping tutorial
nex = WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.XPATH, nextX)))
ActionChains(driver).click(nex).perform()

nex = WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.XPATH, nextX)))
ActionChains(driver).click(nex).perform()

done = WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.XPATH, doneX)))
ActionChains(driver).click(done).perform()

#zooming out to reach copy buttons (Longer answers from ChatGPT will make the copy button unreachable)
pyautogui.keyDown('ctrl')
j: int = 1
for j in range(6):
    pyautogui.press('-')
    time.sleep(1)
pyautogui.keyUp('ctrl')

#Generating a title
msg = WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.XPATH, messageX)))
msg.send_keys('Suggest one (only one) book title with the theme "'+genre+'" for my book that is new (In English)')
send = WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
            (By.XPATH, sendX)))
ActionChains(driver).click(send).perform()
#Clicking the copy button to copy ChatGPT's response
ActionChains(driver).click(WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
            (By.XPATH, copyX)))).perform()
title = ""
#Adding the title as the heading of the document
try:
    title = pyperclip.paste().split('"')[1].replace('"','') #Pasting the response into the title
    document.add_heading(title, 0)
except Exception:
    title = pyperclip.paste().replace('"','')
    document.add_heading(title, 0)

#Generate a random number of chapter titles
ch_number = random.randint(5,12)
msg.send_keys('Now suggest '+str(ch_number)+' chapter titles for the book "'+title+'".')
ActionChains(driver).click(send).perform()
#Clicking the copy button
ActionChains(driver).click(WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.XPATH, copyX+"[2]")))).perform()

#Function to replace all substrings of the pattern "Chapter x:" where x is any number with the string "x. "
def replace_chapter(text):
    #Regular expression pattern to match "Chapter x:" where x is any number
    pattern = r'Chapter (\d+):'
    #Function to process the match and extract the number
    def replace_match(match):
        number = match.group(1)
        return f"{number}. "
    #re.sub() to find all matches and replace them
    result = re.sub(pattern, replace_match, text)
    return result

#Creating a string with the chapter list
try:
    chapter_list = pyperclip.paste().split("1.")[1].replace('"','').replace('Chapter: ', '')
except Exception: #ChatGPT may answer with an unexpected format of the chapter list
    chapter_list = replace_chapter(pyperclip.paste().replace('"',''))
x = 3 #x is the count for the number of responses from ChatGPT. It is used to address the correct copy button

#Writing out the chapter
for i in range(1, ch_number+1):
    #Determine the current chapter for which to write paragraphs
    chapter = ""
    if i == ch_number: #Last chapter -> chapter_list only consists of the title for the last chapter
        chapter = chapter_list
        document.add_heading(chapter, 1)
    else:
        try:
            chapter = chapter_list.split(str(i+1)+".", 1)[0]
            document.add_heading(chapter, 1)
            chapter_list = chapter_list.split(str(i+1)+".", 1)[1] #Removing the current chapter title from the list of remaining chapter titles
        except Exception: #As ChatGPT's responses are sometimes unpredictable exceptions will be handled by using the simplest solution
            chapter = chapter_list
            document.add_heading(chapter, 1)
    p_numbers = 1 + random.randint(2,3) #1 for introduction and random number of subchapters
    #Providing context for ChatGPT will prevent it from getting out of control (Otherwise it sometimes writes chapters for a completely different book, etc.)
    msg.send_keys('Now suggest '+str(p_numbers-1)+' Subchapter-titles for the chapter "'+chapter.replace("\n","")+'" of the book "'+title.replace("\n","")+'".') #Requesting subchapter titles: This will prevent ChatGPT from writing paragraphs with repeating content
    ActionChains(driver).click(send).perform()
    #Scroll down
    try:
        ActionChains(driver).click(WebDriverWait(driver, 5, 1).until(
            expect.visibility_of_element_located(
            (By.XPATH, scrollX)))).perform()
    except Exception:
        pass
    #Clicking the copy button
    ActionChains(driver).click(WebDriverWait(driver, 600, 1).until(
            expect.visibility_of_element_located(
            (By.XPATH, copyX+"["+str(x)+"]")))).perform()
    x+=1 #Incrementing x because the x-th copy button has been clicked
    #Creating a string with the list of subchapters
    subchapters = ""
    try:
        subchapters = pyperclip.paste().split("1.", 1)[1].replace('"','').replace('Subchapter: ', '')
    except Exception:
        subchapters = pyperclip.paste().replace('"','').replace('Subchapter: ', '')
    s="" #String containing the content of the chapter
    pyperclip.copy("") #Resetting the clipboard as it can contain wrong text
    #Writing out the paragraphs
    for j in range(1,p_numbers+1):
        if j==1:
            #Uncomment "continue" and comment out the line after that to skip the introduction of the chapters as it may be unnecessary or ChatGPT may generate a meaningless introduction
            #continue
            msg.send_keys('Write an introduction for the chapter "'+chapter.replace("\n","")+'" of the book "'+title.replace("\n","")+'" (In English)')
        else:
            #Determine the current subchapter which to write
            subchapter = ""
            if j == p_numbers:
                subchapter = subchapters
            else:
                try:
                    subchapter = subchapters.split(str(j)+".", 1)[0]
                    subchapters = subchapters.split(str(j)+".", 1)[1] #Removing the current subchapter title from the list of remaining subchapter titles
                except Exception:
                    subchapter = subchapters
            msg.send_keys('Write a paragraph (containing at least 500 words) for the subchapter "'+subchapter.replace("\n","")+'". It is in the chapter "'+chapter.replace("\n","")+'" of the book "'+title.replace("\n","")+'".')
        ActionChains(driver).click(send).perform()
        #Scroll down
        try:
            ActionChains(driver).click(WebDriverWait(driver, 2, 1).until(
                expect.visibility_of_element_located(
                (By.XPATH, scrollX)))).perform()
        except Exception:
            pass
        #Clicking the copy button
        ActionChains(driver).click(WebDriverWait(driver, 600, 1).until(
            expect.visibility_of_element_located(
            (By.XPATH, copyX+"["+str(x)+"]")))).perform()
        #Scroll down
        try:
            ActionChains(driver).click(WebDriverWait(driver, 5, 1).until(
                expect.visibility_of_element_located(
                (By.XPATH, scrollX)))).perform()
        except Exception:
            pass
        #Pasting the paragraph into the result
        try:
            s+=pyperclip.paste().split("\n", 1)[1].replace('"','').replace('\n', '').rsplit('.',1)[0]+".\n" #rsplit removes last sentence (in case unfinished)
        except Exception:
            try:
                s+=pyperclip.paste().replace('"','').replace('\n', '')+"\n"
            except Exception:
                pass
        x+=1 #Incrementing x because the x-th copy button has been clicked
        pyperclip.copy("")
    document.add_paragraph(s) #Adding the new paragraph to the document

#Saving word-document in directory AutoBooks
directory = "AutoBooks"
if not os.path.exists(directory):
    os.makedirs(directory)
document.save(directory+'/'+title.replace("\n","").replace(":","").replace("/","").replace("?","")+'.docx')

#Generating a description for the book and saving it in a Textfile
msg.send_keys('Now write a description of the book "'+title.replace("\n","")+'" (In English).')
ActionChains(driver).click(send).perform()
ActionChains(driver).click(WebDriverWait(driver, 120, 1).until(
        expect.visibility_of_element_located(
        (By.XPATH, copyX+"["+str(x)+"]")))).perform()
f = open(directory+"/Description-"+title.replace("\n","").replace(":","").replace("/","").replace("?","")+".txt", "a") #Opening Textfile with the name: Description-<simplified title>
#Writing genre and description
f.write(genre+"\n\n")
f.write(pyperclip.paste().replace('\n\n\n', '\n'))
f.close()

#Quitting driver and searching for an element which doesnt exist (because the last element to be accessed over the driver may cause an exception)
driver.quit()
try:
    noExist = WebDriverWait(driver, 1, 1).until(
            expect.visibility_of_element_located(
            (By.ID, "noExist")))
except Exception:
    pass